"""
vula/ingestion/pipeline.py

Vula Document Intelligence Pipeline

Converts client uploads into a searchable, queryable knowledge base.

Flow:
    Upload (PDF/Word/Excel/Image/Scan)
        ↓ GLM-OCR (0.9B local) — image/scan → clean Markdown
        ↓ Docling — structured docs → clean Markdown
        ↓ Chunker — smart semantic splitting
        ↓ BGE-M3 embeddings (local)
        ↓ Qdrant vector store (per-tenant collection)
        ↓ Ready to query via DeepSeek R1

Usage:
    pipeline = VulaIngestionPipeline(tenant_id="abc123")
    results = await pipeline.ingest_file(Path("quote.pdf"))
    print(results.chunks_stored)  # → 42
"""

from __future__ import annotations

import asyncio
import hashlib
import logging
import re
import tempfile
import time
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

import httpx

import sqlite3

from config import settings

logger = logging.getLogger(__name__)


# ─── Ingestion status tracker ─────────────────────────────────────────────────

class IngestionTracker:
    """SQLite log of per-tenant document ingestion status."""

    def __init__(self) -> None:
        self._db = settings.data_dir / "ingestion_log.db"
        self._db.parent.mkdir(parents=True, exist_ok=True)
        self._init()

    def _init(self) -> None:
        with sqlite3.connect(self._db) as conn:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS ingestion_log (
                    id           INTEGER PRIMARY KEY AUTOINCREMENT,
                    tenant_id    TEXT NOT NULL,
                    doc_id       TEXT NOT NULL,
                    filename     TEXT NOT NULL,
                    status       TEXT NOT NULL DEFAULT 'queued',
                    chunks       INTEGER DEFAULT 0,
                    error        TEXT,
                    started_at   TEXT,
                    completed_at TEXT
                )
            """)
            conn.execute("CREATE INDEX IF NOT EXISTS idx_ingest_tenant ON ingestion_log(tenant_id)")
            conn.commit()

    def start(self, tenant_id: str, doc_id: str, filename: str) -> None:
        with sqlite3.connect(self._db) as conn:
            conn.execute(
                "INSERT OR REPLACE INTO ingestion_log (tenant_id, doc_id, filename, status, started_at) VALUES (?,?,?,'processing',datetime('now'))",
                (tenant_id, doc_id, filename),
            )
            conn.commit()

    def complete(self, tenant_id: str, doc_id: str, chunks: int) -> None:
        with sqlite3.connect(self._db) as conn:
            conn.execute(
                "UPDATE ingestion_log SET status='done', chunks=?, completed_at=datetime('now') WHERE tenant_id=? AND doc_id=?",
                (chunks, tenant_id, doc_id),
            )
            conn.commit()

    def fail(self, tenant_id: str, doc_id: str, error: str) -> None:
        with sqlite3.connect(self._db) as conn:
            conn.execute(
                "UPDATE ingestion_log SET status='failed', error=?, completed_at=datetime('now') WHERE tenant_id=? AND doc_id=?",
                (error[:500], tenant_id, doc_id),
            )
            conn.commit()

    def get_all(self, tenant_id: str) -> list:
        with sqlite3.connect(self._db) as conn:
            rows = conn.execute(
                "SELECT doc_id, filename, status, chunks, error, started_at, completed_at FROM ingestion_log WHERE tenant_id=? ORDER BY id DESC",
                (tenant_id,),
            ).fetchall()
        return [
            {"doc_id": r[0], "filename": r[1], "status": r[2], "chunks": r[3],
             "error": r[4], "started_at": r[5], "completed_at": r[6]}
            for r in rows
        ]


_tracker: Optional[IngestionTracker] = None


def get_tracker() -> IngestionTracker:
    global _tracker
    if _tracker is None:
        _tracker = IngestionTracker()
    return _tracker


# ─── Config (from settings / .env) ───────────────────────────────────────────
OLLAMA_BASE = settings.ollama_base
QDRANT_BASE = settings.qdrant_base
GLM_OCR_MODEL = settings.model_ocr
EMBED_MODEL = settings.model_embed
CHUNK_SIZE = settings.chunk_size
CHUNK_OVERLAP = settings.chunk_overlap
MAX_FILE_SIZE_MB = settings.max_file_mb


# ─── Data Classes ─────────────────────────────────────────────────────────────

@dataclass
class DocumentChunk:
    chunk_id: str
    tenant_id: str
    doc_id: str
    filename: str
    page_num: int
    chunk_index: int
    text: str
    embedding: List[float] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)


@dataclass
class IngestionResult:
    doc_id: str
    filename: str
    tenant_id: str
    pages_processed: int
    chunks_stored: int
    file_type: str
    processing_time_s: float
    status: str          # success | partial | failed
    error: Optional[str] = None


# ─── OCR Layer ────────────────────────────────────────────────────────────────

class OCRProcessor:
    """
    Converts scanned images and PDFs to clean Markdown text.

    Primary:  GLM-OCR via Ollama (local, free — but a shared GPU box over a tunnel, so slow
              under load; confirmed via direct reproduction to hit ReadTimeout on image-heavy
              pages even at 180s, which renders as an empty message in older logs).
    Fallback: cloud vision (settings.model_vision, e.g. gemini-2.5-flash) — properly
              provisioned, consistently fast, small per-call cost. Used whenever local OCR
              fails or times out, not just when local is totally unreachable.
    Last resort: pytesseract (pure Python, no GPU/cloud needed) — only if no cloud key is set.
    """

    # 2026-08-17: confirmed live against a real DIGG payment-notification page — GLM-OCR (the
    # small 0.9B local model) didn't fail or time out, it confidently fabricated a completely
    # different, plausible-looking generic business letter (wrong date, wrong amount, and
    # literal "[Name Redacted]"/"[Company Name Redacted]" bracket placeholders) instead of
    # actually reading the image. The old "non-empty text = success" check had no way to catch
    # this, so it never escalated to the far more reliable cloud vision fallback. A genuine OCR
    # read of a real document never contains meta-commentary bracket placeholders like this —
    # only a model that gave up and generated a template instead does.
    _HALLUCINATION_MARKERS = re.compile(
        r"\[(?:name|company|position|amount|date|signature|redacted|unclear|illegible|"
        r"unknown|placeholder)[^\]]{0,40}\]",
        re.IGNORECASE,
    )

    def _looks_hallucinated(self, text: str) -> bool:
        return bool(self._HALLUCINATION_MARKERS.search(text or ""))

    def __init__(self, ollama_base: str = OLLAMA_BASE):
        self.ollama_base = ollama_base

    async def process_image(self, image_path: Path) -> str:
        """Extract text from image/scanned page using GLM-OCR, escalating to cloud vision on
        failure, timeout, OR a hallucinated-looking response, then to local pytesseract as the
        last resort."""
        import base64
        with open(image_path, "rb") as f:
            image_b64 = base64.b64encode(f.read()).decode()

        prompt = "Parse this document to Markdown. Extract all text, tables, and structure accurately."
        payload = {
            "model": GLM_OCR_MODEL,
            "prompt": prompt,
            "images": [image_b64],
            "stream": False,
            "options": {"temperature": 0.1, "num_predict": 4096},
        }

        local_text = ""
        try:
            # The Ollama tunnel is behind Cloudflare Access — send the service-token headers
            # (same ones llm_router uses) or this call is blocked at the edge with a redirect
            # that then fails to parse as JSON.
            from core.llm_router import _ollama_headers
            async with httpx.AsyncClient(timeout=90.0) as client:
                resp = await client.post(
                    f"{self.ollama_base}/api/generate", json=payload,
                    headers=_ollama_headers() or None,
                )
                resp.raise_for_status()
                local_text = resp.json().get("response", "").strip()
                if local_text and not self._looks_hallucinated(local_text):
                    return local_text
                if local_text:
                    logger.warning(f"GLM-OCR output looks hallucinated (bracket placeholders) — "
                                   f"escalating to cloud vision instead of trusting it: {local_text[:200]!r}")
        except Exception as e:
            logger.warning(f"GLM-OCR failed, escalating to cloud vision: {e}")

        # 2026-08-17: cloud vision isn't immune either — reproduced live, the SAME image
        # hallucinated on 2 of 3 consecutive cloud vision calls (identical bracket-placeholder
        # pattern) and only returned the real content on the 3rd. A bounded retry mirrors that —
        # cheap relative to a wrong financial figure being silently trusted.
        cloud_text = ""
        for attempt in range(2):
            cloud_text = await self._cloud_vision_fallback(image_path, prompt)
            if cloud_text and not self._looks_hallucinated(cloud_text):
                return cloud_text
            if cloud_text:
                logger.warning(f"Cloud vision OCR attempt {attempt + 1} looks hallucinated too "
                               f"(bracket placeholders): {cloud_text[:200]!r}")

        tesseract_text = await self._docling_fallback(image_path)
        if tesseract_text:
            return tesseract_text
        # Nothing came back clean — a flagged-but-real response beats an empty page, which
        # upstream treats as a total ingestion failure.
        return cloud_text or local_text or ""

    _IMAGE_MIME = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                   ".webp": "image/webp", ".gif": "image/gif", ".tiff": "image/tiff"}

    async def _cloud_vision_fallback(self, path: Path, prompt: str) -> str:
        """Escalate to the cloud vision model (fast, reliable) — same route used for the
        WhatsApp evidence-photo check. Returns "" if no cloud key is configured."""
        try:
            from core.llm_router import resolve_cloud_vision_route
            route = resolve_cloud_vision_route()
            if not route:
                return ""
            model, api_key, api_base = route

            import base64
            import litellm
            litellm.drop_params = True
            img_b64 = base64.b64encode(path.read_bytes()).decode()
            mime = self._IMAGE_MIME.get(path.suffix.lower(), "image/png")
            resp = await litellm.acompletion(
                model=model,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_b64}"}},
                    ],
                }],
                temperature=0.1,
                max_tokens=4096,
                api_key=api_key,
                api_base=api_base,
            )
            return (resp.choices[0].message.content or "").strip()
        except Exception as e:
            logger.warning(f"Cloud vision OCR fallback failed: {e}")
            return ""

    async def _docling_fallback(self, path: Path) -> str:
        """Pure Python fallback using pdfminer + pytesseract."""
        try:
            import pytesseract
            from PIL import Image
            img = Image.open(path)
            return pytesseract.image_to_string(img)
        except ImportError:
            logger.error("pytesseract not installed — install with: pip install pytesseract pillow")
            return ""


# ─── Document Parser ──────────────────────────────────────────────────────────

class DocumentParser:
    """
    Parses various file types to clean text pages.

    Supported: PDF, DOCX, XLSX, CSV, TXT, MD, PNG, JPG, JPEG, WEBP, GIF
    Placeholder-only (filename/metadata noted, content not read): DWG, DXF
    """

    def __init__(self):
        self.ocr = OCRProcessor()

    async def parse(self, file_path: Path) -> List[tuple[int, str]]:
        """
        Returns list of (page_number, text) tuples.
        Page numbers are 1-indexed.
        """
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return await self._parse_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return await self._parse_docx(file_path)
        elif suffix in (".xlsx", ".xls", ".csv"):
            return await self._parse_spreadsheet(file_path)
        elif suffix in (".png", ".jpg", ".jpeg", ".webp", ".tiff", ".gif"):
            text = await self.ocr.process_image(file_path)
            return [(1, text)] if text else []
        elif suffix in (".txt", ".md"):
            text = file_path.read_text(errors="replace")
            return [(1, text)]
        elif suffix in (".dwg", ".dxf"):
            # No free/native Python reader for Autodesk's binary DWG format (DXF is
            # ASCII/parseable via ezdxf, but not worth a new dependency for one format
            # until it's actually needed). Rather than silently returning nothing —
            # which left every CAD drawing stuck unanalysed with zero information — hand
            # back an honest placeholder so downstream analysis can at least classify it
            # as a drawing from the filename and record that it received one, instead of
            # failing outright. A real content-reading fix needs a DWG→DXF/PDF converter
            # (e.g. the ODA File Converter or a paid conversion API) — not implemented.
            kind = "DWG" if suffix == ".dwg" else "DXF"
            note = (f"[{kind} CAD drawing file — Vula cannot read native {kind} content yet "
                    f"(no converter configured). Filename: {file_path.name}]")
            logger.info(f"{kind} file received, no content reader available: {file_path.name}")
            return [(1, note)]
        else:
            logger.warning(f"Unsupported file type: {suffix}")
            return []

    async def _parse_pdf(self, path: Path) -> List[tuple[int, str]]:
        """Parse PDF — native text extraction first, OCR if scanned or if pdfminer can't parse
        the document at all. Some real-world PDF producers (confirmed 2026-08-17 against a real
        DIGG "Notification of Payment" PDF from FNB) emit non-compliant ASCII85 streams that
        pdfminer's strict decoder rejects outright, raising mid-document. pdfplumber's own
        .to_image() can't help there — it's still built on the same pdfminer object model that
        just failed — so on ANY native-parsing failure this falls back to the whole document via
        pdf2image/poppler (a genuinely independent rendering engine, same pattern already proven
        in vula/takeoff/plan_reader.py's _full_ocr_fallback) rather than failing the document
        outright."""
        try:
            pages = await self._parse_pdf_native(path)
            if pages:
                return pages
        except ImportError:
            logger.warning("pdfplumber not installed: pip install pdfplumber")
        except Exception as exc:
            logger.warning(f"Native PDF parsing failed for {path.name}, falling back to OCR: {exc}")
        return await self._parse_pdf_ocr_fallback(path)

    async def _parse_pdf_native(self, path: Path) -> List[tuple[int, str]]:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                if len(text.strip()) < 50:
                    # Likely scanned — extract page as image and OCR
                    img = page.to_image(resolution=200)
                    img_path = Path(tempfile.gettempdir()) / f"vula_page_{uuid.uuid4().hex}.png"
                    img.save(str(img_path))
                    text = await self.ocr.process_image(img_path)
                    img_path.unlink(missing_ok=True)

                # Also extract tables
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        rows = [" | ".join(str(c) for c in row if c) for row in table if row]
                        text += "\n\n" + "\n".join(rows)

                pages.append((i, text.strip()))
        return [p for p in pages if p[1]]

    async def _parse_pdf_ocr_fallback(self, path: Path) -> List[tuple[int, str]]:
        """Independent-engine (poppler) fallback for PDFs pdfminer can't parse at all — mirrors
        vula/takeoff/plan_reader.py's _full_ocr_fallback."""
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(str(path), dpi=200)
        except Exception as exc:
            logger.error(f"pdf2image/poppler fallback also failed for {path.name}: {exc}")
            # Last resort: treat the whole PDF as one image (matches the old ImportError path).
            return [(1, await self.ocr.process_image(path))]
        pages = []
        for i, img in enumerate(images, 1):
            img_path = Path(tempfile.gettempdir()) / f"vula_pdf_ocr_{uuid.uuid4().hex}.png"
            img.save(str(img_path))
            text = await self.ocr.process_image(img_path)
            img_path.unlink(missing_ok=True)
            pages.append((i, (text or "").strip()))
        return [p for p in pages if p[1]]

    async def _parse_docx(self, path: Path) -> List[tuple[int, str]]:
        """Parse Word document."""
        try:
            from docx import Document
            doc = Document(path)
            full_text = "\n\n".join(
                para.text for para in doc.paragraphs if para.text.strip()
            )
            # Extract tables
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append(" | ".join(cell.text.strip() for cell in row.cells))
                full_text += "\n\n" + "\n".join(rows)
            return [(1, full_text)]
        except ImportError:
            logger.warning("python-docx not installed: pip install python-docx")
            return []

    async def _parse_spreadsheet(self, path: Path) -> List[tuple[int, str]]:
        """Parse Excel/CSV into readable text."""
        try:
            import pandas as pd
            if path.suffix == ".csv":
                df = pd.read_csv(path, dtype=str).fillna("")
            else:
                df = pd.read_excel(path, dtype=str).fillna("")

            # Convert to readable markdown table
            header = " | ".join(df.columns)
            rows = [" | ".join(row) for row in df.values.tolist()]
            text = header + "\n" + "\n".join(rows)
            return [(1, text)]
        except ImportError:
            logger.warning("pandas not installed: pip install pandas openpyxl")
            return []


# ─── Chunker ─────────────────────────────────────────────────────────────────

class SemanticChunker:
    """
    Splits document text into overlapping chunks optimised for RAG.
    
    Strategy: sentence-aware splitting that respects paragraph boundaries,
    preserves table rows, and maintains context via overlap.
    """

    def __init__(self, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        if not text.strip():
            return []

        # Split by double newline (paragraph boundary) first
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        chunks = []
        current_chunk = []
        current_size = 0

        for para in paragraphs:
            para_tokens = len(para.split())

            # If single paragraph exceeds chunk size, split by sentence
            if para_tokens > self.chunk_size:
                sentences = self._split_sentences(para)
                for sent in sentences:
                    sent_tokens = len(sent.split())
                    if current_size + sent_tokens > self.chunk_size and current_chunk:
                        chunks.append(" ".join(current_chunk))
                        # Overlap: keep last N tokens
                        overlap_text = " ".join(current_chunk)
                        overlap_words = overlap_text.split()[-self.overlap:]
                        current_chunk = overlap_words
                        current_size = len(overlap_words)
                    current_chunk.append(sent)
                    current_size += sent_tokens
            else:
                if current_size + para_tokens > self.chunk_size and current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                    overlap_words = "\n\n".join(current_chunk).split()[-self.overlap:]
                    current_chunk = [" ".join(overlap_words)]
                    current_size = len(overlap_words)
                current_chunk.append(para)
                current_size += para_tokens

        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        return [c for c in chunks if len(c.strip()) > 20]

    def _split_sentences(self, text: str) -> List[str]:
        """Basic sentence splitter."""
        import re
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]


# ─── Embedding Layer ──────────────────────────────────────────────────────────

class EmbeddingEngine:
    """
    Generates embeddings using BGE-M3 via Ollama.
    
    BGE-M3: multilingual (works with Afrikaans), 1024-dim, MIT licensed.
    Pull with: ollama pull bge-m3
    """

    def __init__(self, ollama_base: str = OLLAMA_BASE, model: str = EMBED_MODEL):
        self.ollama_base = ollama_base
        self.model = model
        self._dim: Optional[int] = None

    async def embed(self, text: str) -> List[float]:
        """Generate embedding — routed by MODEL_EMBED name, not by provider keys.

        - "text-embedding-*"  → OpenRouter/OpenAI (1536-dim, cloud, always-on)
        - anything else (bge-m3) → Ollama (1024-dim, local/tunnel, free)

        A Qdrant collection has a fixed vector size, so the embed model must
        stay consistent for a given collection. Pick one per environment:
          Cloud-reliable: MODEL_EMBED=text-embedding-3-small
          Free/local:     MODEL_EMBED=bge-m3  (via the vula-ai.com tunnel)
        """
        if self.model.startswith("text-embedding"):
            return await self._embed_openai(text)
        return await self._embed_ollama(text)

    async def _embed_ollama(self, text: str) -> List[float]:
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(
                f"{self.ollama_base}/api/embeddings",
                json={"model": self.model, "prompt": text},
            )
            resp.raise_for_status()
            embedding = resp.json().get("embedding", [])
            if not self._dim and embedding:
                self._dim = len(embedding)
            return embedding

    async def _embed_openai(self, text: str) -> List[float]:
        """OpenAI-compatible embeddings via OpenRouter."""
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(
                "https://openrouter.ai/api/v1/embeddings",
                headers={"Authorization": f"Bearer {settings.openrouter_api_key}"},
                json={"model": self.model, "input": text},
            )
            resp.raise_for_status()
            data = resp.json().get("data", [{}])
            embedding = data[0].get("embedding", []) if data else []
            if not self._dim and embedding:
                self._dim = len(embedding)
            return embedding

    async def embed_batch(self, texts: List[str]) -> List[List[float]]:
        """Embed multiple texts with concurrency limit."""
        semaphore = asyncio.Semaphore(4)  # Max 4 concurrent embedding calls

        async def _embed_one(text: str) -> List[float]:
            async with semaphore:
                return await self.embed(text)

        return await asyncio.gather(*[_embed_one(t) for t in texts])

    @property
    def dimension(self) -> int:
        return self._dim or 1024  # BGE-M3 default


# ─── Vector Store ─────────────────────────────────────────────────────────────

class QdrantStore:
    """
    Per-tenant Qdrant collections for isolated knowledge bases.

    Collection naming: vula_{tenant_id}
    Each client's data is completely isolated.
    Works with local Qdrant and Qdrant Cloud (set QDRANT_API_KEY for cloud).
    """

    def __init__(self, base_url: str = QDRANT_BASE):
        self.base = base_url.rstrip("/")
        self._api_key = settings.qdrant_api_key

    def _headers(self) -> dict:
        if self._api_key:
            return {"api-key": self._api_key}
        return {}

    def _collection_name(self, tenant_id: str) -> str:
        return f"vula_{tenant_id.replace('-', '_')}"

    async def ensure_collection(self, tenant_id: str, vector_size: int = 1024) -> None:
        """Create collection if it doesn't exist, and ensure the source_type payload
        index (required by Qdrant to filter authority-aware retrieval)."""
        name = self._collection_name(tenant_id)
        async with httpx.AsyncClient(timeout=10.0, headers=self._headers()) as client:
            resp = await client.get(f"{self.base}/collections/{name}")
            if resp.status_code != 200:
                resp = await client.put(
                    f"{self.base}/collections/{name}",
                    json={
                        "vectors": {"size": vector_size, "distance": "Cosine"},
                        "optimizers_config": {"default_segment_number": 2},
                    },
                )
                if resp.status_code not in (200, 201):
                    raise RuntimeError(f"Failed to create Qdrant collection: {resp.text}")
                logger.info("Created Qdrant collection: %s", name)
            # Idempotent: ensure the keyword index used by authority filtering.
            try:
                await client.put(
                    f"{self.base}/collections/{name}/index",
                    json={"field_name": "source_type", "field_schema": "keyword"},
                )
            except Exception as exc:
                logger.debug("source_type index ensure skipped for %s: %s", name, exc)

    async def upsert_chunks(self, tenant_id: str, chunks: List[DocumentChunk]) -> int:
        """Store document chunks with embeddings."""
        if not chunks:
            return 0

        name = self._collection_name(tenant_id)
        points = [
            {
                "id": int(hashlib.md5(c.chunk_id.encode()).hexdigest()[:8], 16),
                "vector": c.embedding,
                "payload": {
                    "chunk_id": c.chunk_id,
                    "tenant_id": c.tenant_id,
                    "doc_id": c.doc_id,
                    "filename": c.filename,
                    "page_num": c.page_num,
                    "chunk_index": c.chunk_index,
                    "text": c.text,
                    **c.metadata,
                },
            }
            for c in chunks
            if c.embedding
        ]

        async with httpx.AsyncClient(timeout=30.0, headers=self._headers()) as client:
            resp = await client.put(
                f"{self.base}/collections/{name}/points",
                json={"points": points},
            )
            resp.raise_for_status()

        return len(points)

    async def search(
        self,
        tenant_id: str,
        query_embedding: List[float],
        limit: int = 5,
        score_threshold: float = 0.3,
        exclude_source_types: Optional[List[str]] = None,
        category: Optional[str] = None,
    ) -> List[dict]:
        """Semantic search across tenant's knowledge base.

        `exclude_source_types` filters out low-authority content (e.g. ingested chat
        exports / auto-learned Q&A tagged 'conversation'/'learned') so factual answers
        don't get polluted. Untagged legacy points are never excluded.

        `category` (2026-08-24, structured starter KB): when given, narrows results to
        chunks tagged with that category (see ingest_text's `category` param) — e.g. a
        question that clearly implies "invoice" or "menu" can search just that slice
        instead of ranking across the whole collection. None (default) searches
        everything, unchanged from before this parameter existed.
        """
        name = self._collection_name(tenant_id)
        body: dict = {
            "vector": query_embedding,
            "limit": limit,
            "score_threshold": score_threshold,
            "with_payload": True,
        }
        must_not = [{"key": "source_type", "match": {"value": t}} for t in (exclude_source_types or [])]
        must = [{"key": "category", "match": {"value": category}}] if category else []
        if must_not or must:
            body["filter"] = {k: v for k, v in (("must_not", must_not), ("must", must)) if v}
        async with httpx.AsyncClient(timeout=15.0, headers=self._headers()) as client:
            resp = await client.post(
                f"{self.base}/collections/{name}/points/search",
                json=body,
            )
            if resp.status_code == 404:
                return []
            # If filtering is unavailable (e.g. missing index on a legacy collection),
            # degrade gracefully to an unfiltered search rather than failing the reply.
            if resp.status_code == 400 and "filter" in body:
                logger.warning("Authority filter unavailable for %s — retrying unfiltered", name)
                body.pop("filter", None)
                resp = await client.post(
                    f"{self.base}/collections/{name}/points/search", json=body,
                )
            resp.raise_for_status()
            hits = resp.json().get("result", [])
            results = [hit["payload"] for hit in hits]
            for i, hit in enumerate(hits):
                results[i]["score"] = hit.get("score", 0.0)
            return results

    async def delete_document(self, tenant_id: str, doc_id: str) -> None:
        """Remove all chunks for a specific document."""
        name = self._collection_name(tenant_id)
        async with httpx.AsyncClient(timeout=15.0, headers=self._headers()) as client:
            await client.post(
                f"{self.base}/collections/{name}/points/delete",
                json={"filter": {"must": [{"key": "doc_id", "match": {"value": doc_id}}]}},
            )


# ─── Main Pipeline ────────────────────────────────────────────────────────────

class VulaIngestionPipeline:
    """
    The complete document → knowledge base pipeline for one client tenant.

    Usage:
        pipeline = VulaIngestionPipeline(tenant_id="client_abc")
        result = await pipeline.ingest_file(Path("my_quote_template.pdf"))
        
        # Query after ingestion
        results = await pipeline.query("What do we charge for electrical drawings?")
    """

    def __init__(self, tenant_id: str):
        self.tenant_id = tenant_id
        self.parser = DocumentParser()
        self.chunker = SemanticChunker()
        self.embedder = EmbeddingEngine()
        self.store = QdrantStore()

    async def ingest_file(self, file_path: Path, source_type: str = "document") -> IngestionResult:
        """
        Full pipeline for a single file.
        Returns IngestionResult with stats.
        """
        started = time.time()
        doc_id = self._doc_id(file_path)
        logger.info(f"[{self.tenant_id}] Ingesting: {file_path.name}")
        tracker = get_tracker()
        tracker.start(self.tenant_id, doc_id, file_path.name)

        try:
            # 1. Parse to text pages
            pages = await self.parser.parse(file_path)
            if not pages:
                tracker.fail(self.tenant_id, doc_id, "No text extracted from document")
                return IngestionResult(
                    doc_id=doc_id, filename=file_path.name,
                    tenant_id=self.tenant_id, pages_processed=0,
                    chunks_stored=0, file_type=file_path.suffix,
                    processing_time_s=time.time() - started,
                    status="failed", error="No text extracted from document",
                )

            logger.info(f"[{self.tenant_id}] Parsed {len(pages)} pages from {file_path.name}")

            # 2. Chunk all pages
            all_chunks: List[DocumentChunk] = []
            chunk_index = 0
            for page_num, page_text in pages:
                raw_chunks = self.chunker.chunk(page_text)
                for raw_chunk in raw_chunks:
                    chunk = DocumentChunk(
                        chunk_id=f"{doc_id}_{chunk_index}",
                        tenant_id=self.tenant_id,
                        doc_id=doc_id,
                        filename=file_path.name,
                        page_num=page_num,
                        chunk_index=chunk_index,
                        text=raw_chunk,
                        metadata={
                            "file_type": file_path.suffix,
                            "file_size_kb": file_path.stat().st_size // 1024,
                            "source_type": source_type,
                        },
                    )
                    all_chunks.append(chunk)
                    chunk_index += 1

            logger.info(f"[{self.tenant_id}] Created {len(all_chunks)} chunks")

            # 3. Embed all chunks
            texts = [c.text for c in all_chunks]
            embeddings = await self.embedder.embed_batch(texts)
            for chunk, emb in zip(all_chunks, embeddings):
                chunk.embedding = emb

            # 4. Ensure Qdrant collection exists
            await self.store.ensure_collection(self.tenant_id, self.embedder.dimension)

            # 5. Upsert to Qdrant
            stored = await self.store.upsert_chunks(self.tenant_id, all_chunks)

            elapsed = round(time.time() - started, 2)
            logger.info(f"[{self.tenant_id}] Stored {stored} chunks in {elapsed}s")
            tracker.complete(self.tenant_id, doc_id, stored)

            return IngestionResult(
                doc_id=doc_id,
                filename=file_path.name,
                tenant_id=self.tenant_id,
                pages_processed=len(pages),
                chunks_stored=stored,
                file_type=file_path.suffix,
                processing_time_s=elapsed,
                status="success",
            )

        except Exception as e:
            logger.error(f"[{self.tenant_id}] Ingestion failed for {file_path.name}: {e}")
            tracker.fail(self.tenant_id, doc_id, str(e))
            return IngestionResult(
                doc_id=doc_id, filename=file_path.name,
                tenant_id=self.tenant_id, pages_processed=0,
                chunks_stored=0, file_type=file_path.suffix,
                processing_time_s=time.time() - started,
                status="failed", error=str(e),
            )

    async def ingest_text(self, content: str, filename: str, doc_id: str | None = None,
                          source_type: str = "document", category: str | None = None) -> IngestionResult:
        """Ingest raw text directly — no file needed. Used to seed the training KB.

        `category` (2026-08-24, structured starter KB): an optional doc-type tag (reuses
        vula_filed_documents' existing category vocabulary — see whatsapp.py::_DOC_CATEGORIES)
        stored on every chunk's payload (QdrantStore.upsert_chunks already spreads `metadata`
        into the payload, so no store-layer change was needed) — lets query() narrow retrieval
        to a specific document type when a question clearly implies one. None/omitted behaves
        exactly as before this parameter existed."""
        started = time.time()
        if doc_id is None:
            doc_id = hashlib.md5(f"{self.tenant_id}:{filename}".encode()).hexdigest()[:16]
        try:
            raw_chunks = self.chunker.chunk(content)
            chunk_metadata = {"source": "training_kb", "source_type": source_type}
            if category:
                chunk_metadata["category"] = category
            all_chunks = [
                DocumentChunk(
                    chunk_id=f"{doc_id}_{i}",
                    tenant_id=self.tenant_id,
                    doc_id=doc_id,
                    filename=filename,
                    page_num=1,
                    chunk_index=i,
                    text=chunk,
                    metadata=dict(chunk_metadata),
                )
                for i, chunk in enumerate(raw_chunks)
            ]
            texts = [c.text for c in all_chunks]
            embeddings = await self.embedder.embed_batch(texts)
            for chunk, emb in zip(all_chunks, embeddings):
                chunk.embedding = emb
            await self.store.ensure_collection(self.tenant_id, self.embedder.dimension)
            stored = await self.store.upsert_chunks(self.tenant_id, all_chunks)
            return IngestionResult(
                doc_id=doc_id, filename=filename, tenant_id=self.tenant_id,
                pages_processed=1, chunks_stored=stored, file_type=".md",
                processing_time_s=round(time.time() - started, 2), status="success",
            )
        except Exception as e:
            logger.error(f"[{self.tenant_id}] ingest_text failed for {filename}: {e}")
            return IngestionResult(
                doc_id=doc_id or "", filename=filename, tenant_id=self.tenant_id,
                pages_processed=0, chunks_stored=0, file_type=".md",
                processing_time_s=time.time() - started, status="failed", error=str(e),
            )

    async def ingest_directory(self, dir_path: Path) -> List[IngestionResult]:
        """Ingest all supported files in a directory."""
        supported = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".txt", ".md", ".png", ".jpg", ".jpeg"}
        files = [f for f in dir_path.iterdir() if f.suffix.lower() in supported]
        results = []
        for f in files:
            result = await self.ingest_file(f)
            results.append(result)
        return results

    # Low-authority content excluded from factual retrieval (ingested chat exports
    # and auto-learned Q&A) — see source_type tagging at ingest.
    _NON_AUTHORITATIVE = ["conversation", "learned"]

    async def query(self, question: str, top_k: int = 5,
                    authoritative_only: bool = False, category: str | None = None) -> List[dict]:
        """
        Semantic search across this tenant's knowledge base.
        Returns relevant document chunks for RAG.

        authoritative_only: exclude conversational/learned chunks (chat logs) so
        factual/code answers come only from real documents and references.

        category (2026-08-24): narrow to one document-type slice (see QdrantStore.search) —
        None (default) searches everything, unchanged from before this parameter existed.
        """
        query_embedding = await self.embedder.embed(question)
        exclude = self._NON_AUTHORITATIVE if authoritative_only else None
        return await self.store.search(
            self.tenant_id, query_embedding, limit=top_k,
            score_threshold=0.35 if authoritative_only else 0.3,
            exclude_source_types=exclude, category=category,
        )

    async def answer(
        self,
        question: str,
        context_label: str = "business documents",
        conversation_history: str = "",
    ) -> str:
        """
        Full RAG answer — retrieves context then generates answer via DeepSeek.
        context_label: describes the source in the prompt
        conversation_history: formatted prior exchanges to inject for memory
        """
        # 1. Retrieve relevant chunks
        chunks = await self.query(question)
        if not chunks:
            return "I don't have enough information about that in your business documents yet. Try uploading more documents or rephrase the question."

        from core.prompt_safety import fence, UNTRUSTED_CONTENT_RULE
        context = fence(context_label.upper().replace(" ", "_"), "\n\n---\n\n".join(
            f"[From: {c['filename']}, page {c['page_num']}]\n{c['text']}"
            for c in chunks
        ))

        # 2. Build prompt with optional conversation history
        history_block = (
            f"\nConversation so far:\n{conversation_history}\n"
            if conversation_history else ""
        )

        # 3. Build system + user messages
        system_msg = (
            f"You are Vula, an AI assistant specialising in South African construction and business. "
            f"Answer questions using the {context_label} provided. "
            f"If the answer isn't in the {context_label}, say so clearly. "
            f"Be concise and practical.\n\n" + UNTRUSTED_CONTENT_RULE
        )
        user_msg = (
            f"{context_label.title()}:{context}\n\n"
            f"{history_block}"
            f"Question: {question}\n\nAnswer:"
        )

        # 4. Generate — use litellm which handles Ollama, OpenRouter, and OpenAI uniformly
        import re
        try:
            import litellm
            from core.llm_router import resolve_generation_route
            litellm.drop_params = True

            # Local-first route (Ollama) with OpenRouter fallback when local is down
            model, api_key, api_base = await resolve_generation_route()

            resp = await litellm.acompletion(
                model=model,
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": user_msg},
                ],
                temperature=0.3,
                max_tokens=1024,
                api_key=api_key,
                api_base=api_base,
            )
            raw = resp.choices[0].message.content or ""
        except Exception:
            # Fallback to direct Ollama HTTP if litellm fails
            async with httpx.AsyncClient(timeout=60.0) as client:
                r = await client.post(
                    f"{OLLAMA_BASE}/api/generate",
                    json={"model": settings.model_worker, "prompt": f"{system_msg}\n\n{user_msg}",
                          "stream": False, "options": {"temperature": 0.3, "num_predict": 1024}},
                )
                r.raise_for_status()
                raw = r.json().get("response", "")

        return re.sub(r"<think>.*?</think>", "", raw, flags=re.DOTALL).strip()

    def _doc_id(self, file_path: Path) -> str:
        content = f"{self.tenant_id}:{file_path.name}:{file_path.stat().st_mtime}"
        return hashlib.md5(content.encode()).hexdigest()[:16]


# ─── Quick Test ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys

    async def test():
        if len(sys.argv) < 2:
            print("Usage: python pipeline.py <file_path> [tenant_id]")
            print("Example: python pipeline.py quote.pdf demo_tenant")
            return

        file_path = Path(sys.argv[1])
        tenant_id = sys.argv[2] if len(sys.argv) > 2 else "test_tenant"

        if not file_path.exists():
            print(f"File not found: {file_path}")
            return

        print("\n🌿 Vula Ingestion Pipeline")
        print(f"   Tenant: {tenant_id}")
        print(f"   File:   {file_path.name}")
        print(f"   Size:   {file_path.stat().st_size // 1024} KB\n")

        pipeline = VulaIngestionPipeline(tenant_id=tenant_id)
        result = await pipeline.ingest_file(file_path)

        print(f"✓ Status:  {result.status}")
        print(f"✓ Pages:   {result.pages_processed}")
        print(f"✓ Chunks:  {result.chunks_stored}")
        print(f"✓ Time:    {result.processing_time_s}s")

        if result.status == "success":
            print("\n💬 Test query...")
            answer = await pipeline.answer("What is this document about?")
            print(f"Answer: {answer[:300]}...")

    asyncio.run(test())
